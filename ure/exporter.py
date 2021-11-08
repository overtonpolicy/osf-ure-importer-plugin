import os,sys
import re
from warnings import warn
import pprint

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt, Cm, RGBColor

# PIL is necessary to get the image size.
from PIL import Image 

# request are needed to fetch images based on their source links.
# they are saved to a tempfile in order to be loaded
import tempfile
import requests

# etree and latex2maathml are necessary to process math equations
from lxml import etree
import latex2mathml.converter

# mistune translates markdown into an Abstract Syntax Tree (AST), 
# which is what we process to generate the export
import mistune

class BaseExporter():

    def __init__(self, include_components=True, wiki_break_type=None, component_break_type='page', add_component_titles=True, add_wiki_titles=False):
        self.add_component_titles = add_component_titles        
        self.add_wiki_titles = add_wiki_titles        
        self.include_components = include_components
        self.wiki_break_type = wiki_break_type
        self.component_break_type = component_break_type

    def process_markdown(self, markdown):
        raise Exception("process_source must be defined in the subclass to return a ure-markdown datastructure")


class Docx(BaseExporter):

    def __init__(self, *args, style_template=None, **kwargs):
        super().__init__(*args, **kwargs)
        if not style_template:
            raise Exception("Need a default style template")
        else:
            self.style_template = style_template
        if not os.path.exists(self.style_template):
            raise Exception(f"{self.style_template} is not the path to a file that can be used for a template")

        # set up the renderer
        self.markdown_analyzer = mistune.create_markdown(
            renderer='ast',
            plugins=[
                'table', 
                'strikethrough', 
                'url',
                'math' 
            ]
        )

    def process_markdown(self, osfmarkdown):
        doc = docx.Document(os.path.abspath(self.style_template))

        for icomp, (title, component) in enumerate(osfmarkdown):
            if icomp > 0:
                # render a break
                self.render_break(doc, self.component_break_type)            
                # every component after the root project
                if self.add_component_titles:
                    doc.add_heading(title, level=1)
            for iwiki, (wikititle, content) in enumerate(component):
                # every wiki after the home 
                if iwiki > 0:
                    self.render_break(doc, self.wiki_break_type)  
                    if self.add_wiki_titles:
                        doc.add_heading(wikititle, level=1)          
                self.render_section(doc, content)
        
        # now save to the output 
        #doc.save(os.path.abspath(args.output))        
        return(doc)


    #
    # Blocks
    #
    def process_newline(self, doc, element):
        return(doc.add_paragraph("\n"))

    def process_thematic_break(self, doc, element):
        # -- We already handle breaks separately, so we don't do anything here.
        # NOTE: This may not be intended
        #return(self.render_break(doc, 'page'))
        return([])

    def process_block_code(self, doc, element):
        if 'info' in element:
            warn(f"Block_Code may have info param: {element['info']}, which is not processed yet")
        return(doc.add_paragraph(
            text=element['text'],
            style="Code Block"
        ))

    def process_heading(self, doc, element):
        head = doc.add_heading(

            level = element['level'],
        )      
        for child in element['children']:
            self.render_element(head, child) 
        return(head)

    def process_block_quote(self, doc, element, style=None):
        """style, if provided, will be ignored """
        children = []
        for i, child in enumerate(element['children']):
            if child['type'] not in ('paragraph', 'block_text'):
                raise Exception(f"Unknown type as child of block quote '{child['type']}'")
            if i == 0: 
                para = doc.add_paragraph(style="Quote")
            else:
                para = doc.add_paragraph(style="Indented Quote")
            for subelem in child['children']:
                self.render_element(para, subelem)
            children.append(para)
        return(children)

    def process_list(self, doc, element):
        level = element['level']
        is_ordered = element['ordered']
        #import pdb; pdb.set_trace()
        children = []
        for i, child in enumerate(element['children']):
            children.append(self.render_list_item(doc, child, ordered=is_ordered, level=level))
        return(children)

    def render_list_item(self, doc, element, ordered, level):
        item_level = element['level']
        #import pdb; pdb.set_trace()
        if ordered:
            list_style = "List Number" if item_level == 1 else f"List Number {item_level}"
        else:
            list_style = "List Bullet" if item_level == 1 else f"List Bullet {item_level}"

        children = []
        for child in element['children']:
            if child['type'] in ('paragraph', 'block_text'):
                xren = self.render_element(doc, child, style=list_style)
            else:
                xren = self.render_element(doc, child)

            if type(xren) is list:
                children.extend(xren)
            else:
                children.append(xren)
        return(children)

    def process_block_html(self, doc, element):
        warn("Found html bock as follows. Delegating to code block:\n{element['text']}")
        return(doc.add_paragraph(
            text=element['raw'],
            style="Code"
        ))

    def process_block_text(self, *args, **kwargs):
        return(self.process_paragraph(*args, **kwargs))

    def process_paragraph(self, doc, element, style=None):
        para = doc.add_paragraph(style=style)
        children = []
        for child in element['children']:
            children.append(self.render_element(para, child))
        return(children)

    def process_mathblock(self, container, element):
        warn("Redirecting analysis of mathblock to codeblock")
        word_math = latex_to_word(element['expression'])

        children = []
        if isinstance(container, docx.text.paragraph.Paragraph):
            # a little hacky, you can't add a paragraph to a paragraph, so we need to add a run with breaks
            run = container.add_run()
            run.add_break()
            children.append(run)
            container._element.append(word_math)
            #children.append(self.process_codespan(container, {'text': element['expression']}))
            run = container.add_run()
            run.add_break()
            children.append(run)
            return(children)
        else:
            para = container.add_paragraph()
            para._element.append(word_math)
            return(para)

        return(self.process_block_code(container, {'text': element['expression']}))

    def process_table(self, doc, element):
        # add table ------------------
        row_count = 0
        col_count = 0
        table_data = []
        for segment in element['children']:
            if segment['type'] == 'table_head':
                # -- table head is a row
                row_count += 1
                if len(segment['children']) > col_count:
                    col_count = len(segment['children'])
                table_data.append(segment)
            else:
                for row in segment['children']:
                    row_count += 1
                    if len(row['children']) > col_count:
                        col_count = len(row['children'])
                    table_data.append(row)  

        table = doc.add_table(rows=row_count, cols=col_count)
        table.style = 'Plain Table 3'

        for irow, row in enumerate(table_data):
            if row['type'] not in ('table_row', 'table_head'):
                raise Exception(f"Expected table row or head but got '{row['type']}'")
            for icol, elem in enumerate(row['children']):
                if elem['type'] not in ('table_cell'):
                    raise Exception(f"Expected table cell  but got '{elem['type']}'")
                cell = table.cell(irow, icol)
                if elem['is_head']:
                    style = 'Table Heading'
                else:
                    style = 'Table Contents'    
                para = cell.paragraphs[0]
                para.style = style
                for child in elem['children']:
                    self.render_element(para, child)            
        return(table)

    #
    # Inline 
    #
    def process_text(self, container, element):
        return(container.add_run(element['text']))

    def process_link(self, container, element):

        result = []
        if not isinstance(container, docx.text.paragraph.Paragraph):
            container = container.add_paragraph()
            result.append(container)

        # This gets access to the document.xml.rels file and gets a new relation id value
        part = container.part
        r_id = part.relate_to(
            element['link'], 
            docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, 
            is_external=True,
        )
        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
        
        children = element['children'] or [{'type': 'text', 'text': element['link']}] 
        child_runs = []
        while children:
            child = children.pop(0)
            # there are weird cases when mistune renders links children of links.
            # If that's the case we're just going add them to the children list
            if child['type'] == 'link':
                if child['children']:
                    children.extend(child['children'])
                continue

            run = self.render_element(container, child)
            if type(run) is list:
                child_runs.extend(run)
            else:
                child_runs.append(run)
        
        for run in child_runs:
            # set the link color
            run.font.color.rgb = RGBColor.from_string('000080')
            # underline the link
            run.underline = True

            xml_rPr = docx.oxml.shared.OxmlElement('w:rPr')
            run._r.append(xml_rPr)
            hyperlink.append(run._r)

        container._p.append(hyperlink)
        return(container)

    def process_image(self, container, element):
        resp = requests.get(element['src'], stream=True)
        if not resp.ok:
            raise Exception(f"Could not download image at {element['src']}")

        extm = re.search(r'\.([^\.]+)$', element['src'].strip())
        if extm:
            ext = extm.group(1)
        else:
            ext = '.png' # just guessing
        img = None
        with tempfile.NamedTemporaryFile(suffix=ext) as fh:
            for chunk in resp:
                fh.write(chunk)
            fh.flush()
            fh.seek(0)

            #read the image
            image = Image.open(fh.name)

            #image size - we will limit all images to the lesser of 5" wide or tall at 72 DPI
            width = image.size[0]
            height = image.size[1]

            max_width = max_height = None
            if width > 72 * 5:
                max_width = docx.shared.Inches(5)
            if height > 72 * 5:
                if not max_width or height > width:
                    max_height = docx.shared.Inches(5)
                    max_width = None

            print(f"LIMITED to {max_width} x {max_height}")
            if isinstance(container, docx.text.paragraph.Paragraph):
                # paragraphs cannot add pictures, but runs can.
                img = (container.add_run().add_picture(fh, width=max_width, height=max_height))
            else:
                img (container.add_picture(fh, width=max_width, height=max_height))
        return(img)

    def process_strikethrough(self, container, element): 
        
        runs = []
        for child in element['children']:
            run = self.render_element(container, child)
            if type(run) is list:
                for r in run:
                    r.font.strike = True
                runs.extend(run)
            else:
                run.font.strike = True
                runs.append(run)
        return(runs)

    def process_emphasis(self, container, element): 
        #italic
        runs = []
        for child in element['children']:
            run = self.render_element(container, child)
            if type(run) is list:
                for r in run:
                    r.italic = True
                runs.extend(run)
            else:
                run.italic = True
                runs.append(run)
        return(runs)

    def process_strong(self, container, element): 
        # bold
        runs = []
        for child in element['children']:
            run = self.render_element(container, child)
            if type(run) is list:
                for r in run:
                    r.bold = True
                runs.extend(run)
            else:
                run.bold = True
                runs.append(run)
        return(runs)

    def process_linebreak(self, container, element):
        return(container.add_run("[linebreak]\n"))

    def process_inline_html(self, container, element):
        # delegete to codespan
        warn(f"designating inline html '{element['text']}' as codespan")
        return(self.process_codespan(container, element))

    def process_codespan(self, container, element):
        return(container.add_run(element['text'], style="Code"))

    def process_mathspan(self, container, element):
        warn("Redirecting analysis of mathspan to codespan")
        word_math = latex_to_word(element['expression'])
        return(container._element.append(word_math))
        #return(self.process_codespan(container, {'text': element['expression']}))

    def render_break(self, doc, break_type):
        if not break_type:
            return
        if break_type == 'page':
            doc.add_page_break()
        elif break_type == 'section':
            raise Exception("New Section not yet handled")
            section = doc.add_section()
        elif break_type == 'line':
            self.process_newline(doc, None)
        else:
            raise Exception(f"Unknown break '{break_type}'")

    def render_element(self, docelem, element, **kwargs):
        funcname = 'process_' + element['type']
        func = getattr(self, funcname)
        return(func(docelem, element, **kwargs))

    def render_section(self, doc, mkdtext):
        ast = self.markdown_analyzer(mkdtext)
        stringified = pprint.pformat(ast, width=200, indent=5) 
        with open('current_mkd.md', 'w') as fh:
            fh.write(mkdtext)
        with open('current_ast.js', 'w') as fh:
            fh.write(stringified)
        
        print(stringified)
        for element in ast:
            self.render_element(doc, element)
        


def latex_to_word(latex_input):
    mathml = latex2mathml.converter.convert(latex_input)
    tree = etree.fromstring(mathml)
    xslt = etree.parse(
        'conf/MML2OMML.XSL'
        )
    transform = etree.XSLT(xslt)
    new_dom = transform(tree)
    return new_dom.getroot()
